# -*- Product under GNU GPL v3 -*-
# -*- Author: E.Aivayan -*-
import glob
import logging
import os
import platform
import re
import subprocess
import tempfile
from pathlib import Path
from shutil import copyfile
from typing import Tuple, Union

from behave.parser import parse_file
from docx import Document
from htmldocx import HtmlToDocx
from markdown_it import MarkdownIt
from matplotlib import pyplot as plt
from PIL import Image

log = logging.getLogger(__name__)


class ExportUtilities:

    def __init__(self, feature_repository: str = None,
                 user_story_tag_prefix: str = None,
                 report_title: str = None):
        self.__feature_repository = feature_repository
        self.__user_story_tag_prefix = user_story_tag_prefix
        self.__report_title = report_title
        self.__document = None
        try:
            version = subprocess.check_output(['java', '-version'], stderr=subprocess.STDOUT)
        except Exception as exception:
            log.warning(f"Check java version error\n JRE might not be installed."
                        f"\n Subprocess spawn {exception.args}")
            version = ""

        self.__jre_present = bool(version)
        self.__jar_path = f"{os.path.dirname(os.path.realpath(__file__))}/assets/plantuml.jar"
        self.__current_feature_tags = None
        self.__include_result = False
        self.__forewords_folder = None
        self.__inline_counter = 0

    @property
    def feature_repository(self):
        return self.__feature_repository

    @feature_repository.setter
    def feature_repository(self, repository: str):
        if Path(repository).exists() and Path(repository).is_dir():
            self.__feature_repository = repository
        else:
            raise FileExistsError(f"{repository} is not a existing folder")

    @property
    def report_title(self):
        return self.__report_title

    @report_title.setter
    def report_title(self, title: str):
        if isinstance(title, str) and title:
            self.__report_title = title
        else:
            raise AttributeError(f"{title} must be a non empty string")

    @property
    def us_tag(self):
        return self.__user_story_tag_prefix

    @us_tag.setter
    def us_tag(self, tag_value: str):
        if isinstance(tag_value, str) and tag_value:
            self.__user_story_tag_prefix = tag_value
        else:
            log.warning("Use the None value")
            self.__user_story_tag_prefix = None

    @property
    def forewords_folder(self) -> Union[str, Path, None]:
        return self.__forewords_folder

    @forewords_folder.setter
    def forewords_folder(self, folder: str):
        if Path(folder).exists() and Path(folder).is_dir():
            self.__forewords_folder = folder
            self.__include_result = True
        else:
            raise FileExistsError(f"{folder} is not a existing folder")

    @property
    def document(self):
        return self.__document

    def _get_level(self, level_name: str) -> int:
        """Return the effective level depending on the inclusion (asis or +1)"""
        level = {"h1": 1,
                 "h2": 2,
                 "h3": 3,
                 "h4": 4,
                 "h5": 5}
        return level[level_name] + 1 if self.__include_result else level[level_name]

    def __forewords_schema_replacement(self, match_obj):
        """Replace schema tags with the generated schema picture"""
        generated_path = self.__generate_diagrams(match_obj.group(1), False)
        generated_path = re.sub(r'\\', '/', generated_path)
        return f"\n![Schema]({generated_path})\n"

    def __forewords_inline_puml(self, match_obj):
        """Generate inline puml and insert"""
        if self.__jre_present:
            current = self.__inline_counter
            temp_puml = Path(f"{tempfile.gettempdir()}/inline_p_{current}.puml")
            temp_puml = temp_puml.resolve()
            with open(temp_puml, "w") as file:
                file.write(match_obj.group(1))
            subprocess.run(
                ["java", "-Djava.awt.headless=true",
                 "-jar", Path(self.__jar_path).absolute(),
                 temp_puml,
                 "-o", Path(f"{tempfile.gettempdir()}")])
            gen_pic_path = Path(f'{tempfile.gettempdir()}/{temp_puml.name.split(".")[0]}.png')
            self.__resize_schema(gen_pic_path)
            generated_path = re.sub(r'\\',
                                    '/',
                                    str(gen_pic_path.absolute()))
            self.__inline_counter += 1
            return f"\n![Diag {current}]({generated_path})\n"
        else:
            return ""

    def __forewords_picture(self, match_obj):
        generated_path = re.sub(
            r'\\',
            '/',
            str(Path(f"{self.forewords_folder}/{match_obj.group(2)}").absolute()))
        self.__resize_schema(Path(generated_path))
        return f"\n![{match_obj.group(1)}]({generated_path})\n"

    def create_application_documentation(self, report_file=None, output_file_name="demo.docx"):
        """
        Create a document (docx) object and read first all ".feature" files and
        add their contents into the document.

        If a report file name is provided, it will add a "last execution" section containing
        the data found in the report.

        The report file must be the "plain" report output file generated from behave.

        :param report_file: The report file path (absolute or relative)
        :param output_file_name : The exported file name by default "demo.docx"
        :return: None
        """
        log.info("Start application documentation")

        # Check which drive is it if window
        self.__document = Document()
        self.document.add_heading(f"{self.__report_title}", 0)  # Document title
        self.document.add_page_break()
        # Is copying file needed
        if platform.system() == "Windows":
            current_execution = os.path.abspath(os.getcwd())
            repository = os.path.abspath(self.__feature_repository)
            if current_execution.split(os.path.sep)[0] == repository.split(os.path.sep)[0]:
                please_copy = False
            else:
                please_copy = True
                temp_file = os.path.abspath(f"{current_execution}/temp.feature")
        else:
            please_copy = False

        log.info(f"We are on {platform.system()} and we need to copy is set to {please_copy}")

        # Insert forewords sections
        if self.forewords_folder is not None:
            self.document.add_heading("Forewords", 1)
            list_of_files = sorted(filter(os.path.isfile,
                                          glob.glob(f"{self.forewords_folder}/*.md")))
            for file in list_of_files:
                with open(file) as foreword_section:
                    content = foreword_section.read()
                    # Shift title level +1
                    content = re.sub(r'^(#*)', r'#\1', content)
                    # Process included picture with relative path
                    content = re.sub(r'!\[([^\[\]]+)\]\(([^\s]+)\)',
                                     self.__forewords_picture,
                                     content)
                    if self.__jre_present:
                        # Process inline puml diagrams
                        content = re.sub(r'```puml[\r|\n]{1,2}([^`]*)```',
                                         self.__forewords_inline_puml,
                                         content,
                                         flags=re.MULTILINE)
                        # Process diagrams
                        content = re.sub(r'!!Workflow:\s*([\.\d\w\-\_\\\/]*)\s*',
                                         self.__forewords_schema_replacement,
                                         content)
                    insert_text(self.document, content)
            self.document.add_page_break()

        if report_file is not None or self.__include_result:
            self.__include_result = True
            self.document.add_heading("Living documentation", 1)

        for file in glob.iglob(f"{self.__feature_repository}/**/*.feature",
                               recursive=True):  # Use the iterator as it's cleaner
            # log.debug()
            log.info(f"Computing {os.path.abspath(file)}")
            if please_copy:
                copyfile(os.path.abspath(file), temp_file)
            else:
                temp_file = os.path.abspath(file)
            try:
                # if file drive is not current cp file and use that copy
                # Use the Behave parser in order to read the feature file
                test = parse_file(temp_file)

                self.add_heading(feature=test)
                self.add_description(feature=test)
                self.add_background(feature=test)
                self.add_scenario(feature=test)
                self.document.add_page_break()
                # rm the file copy
                if please_copy:
                    os.remove(temp_file)
            except Exception as exception:
                log.error(exception)

        if report_file is not None:
            self.add_report(file=report_file)
        if not output_file_name.endswith(".docx"):
            self.document.save(f"{output_file_name}.docx")
        else:
            self.document.save(output_file_name)
        log.info("Processing done.")

    def add_heading(self, feature=None):
        """
        Add the feature name as top level section
        :param feature: the feature object
        :return:
        """
        try:
            log.info(f"Processing {feature.name}")
            self.document.add_heading(feature.name, self._get_level("h1"))
            # TODO work on condition so that no orphan empty paragraph
            paragraph = self.document.add_paragraph("")
            tagged = False
            if self.us_tag is not None:
                matcher = [elem for elem in feature.tags if self.us_tag in elem]
                if matcher:
                    run = paragraph.add_run()
                    run.underline = True
                    run.text = "Related to the user story: "
                    paragraph.add_run(f"{str(matcher).strip('[]')}")
                    tagged = True
            if feature.tags:
                if tagged:
                    paragraph = self.document.add_paragraph("")
                paragraph.add_run(f"Feature tags are {str(feature.tags).strip('[]')}")
        except Exception as exception:
            log.error(exception)
            raise Exception(exception) from exception

    def __generate_diagrams(self, diagram_path,
                            is_feature_diagram: bool = True) -> Union[str, None]:
        """Generate the diagram using plantuml. Return the picture absolute path"""
        try:
            base_path = self.feature_repository if is_feature_diagram else self.forewords_folder
            # Assuming that the diagram path is relative to feature folder repository
            path = Path(f"{base_path}/{diagram_path}")
            resolved = path.resolve()
            # Check if existing png files exists
            if (Path(f"{resolved.name.split('.')[0]}.png").exists()
                    and Path(f"{resolved.name.split('.')[0]}.png").is_file()):
                gen_pic_path = Path(f"{resolved.name.split('.')[0]}.png")
            else:
                # Generate the temporary picture
                try:
                    tmp_pic_folder = Path(f"{tempfile.gettempdir()}")
                    subprocess.run(
                        ["java", "-Djava.awt.headless=true",
                         "-jar", Path(self.__jar_path).absolute(),
                         resolved,
                         "-o", tmp_pic_folder])
                    gen_pic_path = Path(f'{tmp_pic_folder}/{resolved.name.split(".")[0]}.png')
                except FileNotFoundError as file_not_found:
                    # Don't break the flow
                    log.warning(file_not_found.args[0])
                    return

            self.__resize_schema(gen_pic_path)
            return str(gen_pic_path.absolute())
        except Exception as exception:
            log.error(exception)
            raise Exception(exception) from exception

    @staticmethod
    def __resize_schema(schema_picture_path: Path):
        # Resize the picture if too big
        image: Image = Image.open(schema_picture_path)
        width, height = image.size
        # Preserve image ratio
        ratio = width / height
        # TODO resize image using PIL and save to gen_pic_path
        if width > 580 and height < 841:
            new_width = 580
            new_height = int(580 / ratio)
            # image.resize()
            # self.document.add_picture(str(gen_pic_path.absolute()),
            #                           width=Cm(15),
            #                           height=Cm(15 / ratio))
        elif width < 580 and height > 841:
            new_height = 841
            new_width = int(841 * ratio)
            # self.document.add_picture(str(gen_pic_path.absolute()),
            #                           width=Cm(20 * ratio),
            #                           height=Cm(20))
        elif width > 580 and height > 841:
            reduce_factor = max(width / 580, height / 841)
            new_width = int(
                (width * 580) / (580 * reduce_factor))  # Double proportionality on ratio
            # and pixel against cm
            new_height = int((height * 841) / (841 * reduce_factor))
            # self.document.add_picture(str(gen_pic_path.absolute()),
            #                           width=Cm(new_width),
            #                           height=Cm(new_height))
        else:
            new_width = width
            new_height = height
            # self.document.add_picture(str(gen_pic_path.absolute()))
        # TODO return path so that it can be included in generator
        new_image = image.resize((new_width, new_height))
        new_image.save(str(schema_picture_path.absolute()), format="png")

    def __schema_replacement(self, match_obj):
        result = f"!!Workflow: {match_obj.group(1)}\n"
        if self.__jre_present:
            generated_path = self.__generate_diagrams(match_obj.group(1))
            generated_path = re.sub(r'\\', '/', generated_path)
            result = (f"\n![Schema]({generated_path})\n"
                f"!!Workflow: {match_obj.group(1)}\n")

        return result

    def add_description(self, feature=None):
        """
        Add the feature description into the document.
        If the feature file line contains a "*" character it will create a bullet list.
        If the line contains the sentence "business rules" then it will print it in bold.
        Otherwise it will print the line as-is.
        :param feature: the feature object
        :return: None
        """
        try:
            # TODO consider feature.description as markdown.
            # TODO make relevant transformation for Business Rules, workflow and user story
            description = "\n".join(feature.description)
            # Capture user story and format

            description = re.sub(r"as([ \w\"'\.\-]*)",
                                 lambda x:f"<b>As</b> {x.group(1)} <br />",
                                 description,
                                 flags=re.IGNORECASE)
            description = re.sub(r"i want([ \w\"'\.\-]*)",
                                 lambda x: f"<b>I want</b> {x.group(1)} <br />",
                                 description,
                                 flags=re.IGNORECASE)
            description = re.sub(r"so that([ \w\"'\.\-]*)",
                                 lambda x: f"<b>So that</b> {x.group(1)} <br />",
                                 description,
                                 flags=re.IGNORECASE)
            # replace business rules with title h2
            description = re.sub(r'[Bb]usiness [Rr]ules.*',
                                 f"{'#' * self._get_level('h2')} Business Rules",
                                 description)
            # replace !!Workflow: tag with the generated picture inclusion
            description = re.sub(r'!!Workflow:\s*([\.\d\w\-\_\\\/]*)\s*',
                                 self.__schema_replacement,
                                 description)
            # include md description in the document
            insert_text(self.document, description)
            # for line in feature.description:
            #     if re.match(r'\*.*', line):
            #         self.document.add_paragraph(line[1:], style='List Bullet')
            #     elif re.match('[Bb]usiness [Rr]ules.*', line):
            #         paragraph = self.document.add_paragraph("")
            #         paragraph.add_run(line).bold = True
            #     elif re.match(r'!!Workflow:.*', line) and self.__jre_present:
            #         match = re.match(r'^!!Workflow:\s*([\.\d\w\-\_\\\/]*)\s*$', line)
            #         if match:
            #             self.document.add_picture(self.__generate_diagrams(match.group(1)))
            #         self.document.add_paragraph(line)
            #     else:
            #         self.document.add_paragraph(line)
        except Exception as exception:
            log.error(exception)
            raise Exception(exception) from exception

    def add_background(self, feature=None):
        """
        Add the background of each scenario into the document
        :param feature: the feature object from where to retrieve the background
        :return: None
        """
        try:
            if feature.background is not None:
                self.print_scenario_title(scenario_keyword=feature.background.keyword,
                                          scenario_name=feature.background.name)
                self.print_steps(steps=feature.background.steps)
        except Exception as exception:
            log.error(exception)
            raise Exception(exception) from exception

    def add_scenario(self, feature=None):
        """
        Add in the document all the scenarios attached to a feature.
        :param feature: the feature object
        :return: None
        """
        try:
            if feature.scenarios is not None:
                for scenario in feature.scenarios:
                    log.info(f"Processing scenario {scenario.name}")
                    self.print_scenario_title(scenario_keyword=scenario.keyword,
                                              scenario_name=scenario.name,
                                              level=self._get_level("h2"))
                    paragraph = self.document.add_paragraph("Scenario tags are ",
                                                            style='No Spacing')
                    if feature.tags:
                        paragraph.add_run(", ".join({f"'{tag}'" for tag in feature.tags}))
                    if scenario.tags:
                        paragraph.add_run(", ")
                        paragraph.add_run(", ".join({f"'{tag}'" for tag in scenario.tags}))
                    self.print_steps(steps=scenario.steps)
                    if scenario.type == 'scenario_outline':
                        self.print_examples(examples=scenario.examples)
        except Exception as exception:
            log.error(exception)
            raise Exception(exception) from exception

    def print_examples(self, examples=None):
        """
        Add an example section for each example attached to a scenario outline
        :param examples: an example list
        :return: None
        """
        try:
            for example in examples:
                log.info(f"Processing example '{example.name}'")
                self.print_scenario_title(scenario_keyword=example.keyword,
                                          scenario_name=example.name,
                                          level=self._get_level("h3"))
                self.print_table(table=example.table)
        except Exception as exception:
            log.error(exception)
            raise Exception(exception) from exception

    def print_scenario_title(self,
                             scenario_keyword=None,
                             scenario_name=None,
                             level=2):
        """
        Print a section title in the document in the format keyword : name with a level 2
        :param scenario_keyword: the keyword such as Scenario or Scenario Outline or Example
        :param scenario_name: the scenario name
        :param level: the section level. By default level 2
        :return: None
        """
        self.document.add_heading(f"{scenario_keyword}: {scenario_name}", level=level)
        # paragraph = document.add_paragraph("")
        # paragraph.add_run("{}:".format(scenario_keyword)).bold = True
        # paragraph.add_run(scenario_name)

    def print_steps(self, steps=None):
        """
        Add a step section into the document.
        :param steps: the feature steps table
        :return: None
        """
        step_done = []
        for step in steps:
            if step.keyword in step_done:
                keyword = "And"
            else:
                step_done.append(step.keyword)
                keyword = step.keyword

            paragraph = self.document.add_paragraph("", style='No Spacing')
            paragraph.add_run(keyword).bold = True
            paragraph.add_run(" {}".format(step.name))

            if step.table is not None:
                self.print_table(table=step.table)

    def print_table(self, table=None):
        """
        Add a tabular word object in the document based on the feature table object
        :param table: a feature table object
        :return: None
        """
        number_of_column = len(table.headings)

        table_instance = self.document.add_table(rows=1,
                                                 cols=number_of_column,
                                                 style='Light List Accent 3')
        header_cells = table_instance.rows[0].cells
        for count, text in enumerate(table.headings):
            header_cells[count].text = str(text)

        for row in table.rows:
            row_cells = table_instance.add_row().cells
            for count, cell in enumerate(row.cells):
                row_cells[count].text = str(cell)

    def add_report(self, file=None):
        """
        Add a last execution section to a document. It reads an execution plain file report
        :param file: the file name (relative or absolute) of the execution report.
        :return: None
        """
        self.document.add_heading("Last Execution report", 1)
        reporter = {}
        total, succeed, failed = self.__parse_report(file, reporter)

        self.document.add_page_break()
        self.document.add_heading("Last Execution summary", 1)

        labels = ["succeed", "failed", "skipped"]
        part_succeed = int(100 * succeed / total)
        part_failed = int(100 * failed / total)
        sizes = [part_succeed, part_failed, 100 - part_succeed - part_failed]
        fig1, ax1 = plt.subplots()
        ax1.pie(sizes,
                labels=labels,
                colors=['tab:green', 'tab:red', 'tab:gray'],
                autopct='%1.1f%%',
                shadow=True)
        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

        tmp_pic_folder = Path(f"{tempfile.gettempdir()}/result.png")
        plt.savefig(str(tmp_pic_folder.absolute()))
        plt.close(fig1)

        self.document.add_picture(str(tmp_pic_folder.absolute()))

        table_instance = self.document.add_table(rows=1, cols=3, style='Light List Accent 3')
        header_cells = table_instance.rows[0].cells
        header_cells[0].text = "Feature"
        header_cells[1].text = "Scenario"
        header_cells[2].text = "Status"
        log.debug(reporter)
        feature_keys = sorted(reporter.keys())
        for feature_key in feature_keys:
            scenario_keys = sorted(reporter[feature_key].keys())
            log.debug(f"Create table summary {scenario_keys}")
            for scenario_key in scenario_keys:
                row_cells = table_instance.add_row().cells
                row_cells[0].text = feature_key
                if scenario_key is None:
                    row_cells[2].text = "No scenario to report"
                    continue
                row_cells[1].text = scenario_key
                row_cells[2].text = reporter[feature_key][scenario_key]

    def __parse_report(self, file: str = None, reporter: dict = None) -> Tuple[int, int, int]:
        current_feature = None
        current_scenario = None
        last_status = "skipped"
        with open(file) as report_file:
            failed_count = 0
            succeed_count = 0
            test_count = 0
            for line in report_file.readlines():
                if re.match("Feature.*", line):
                    if current_feature is not None:
                        reporter[current_feature].update({current_scenario: last_status})
                        if last_status == "failed":
                            failed_count += 1
                        elif last_status == "passed":
                            succeed_count += 1
                        test_count += 1
                        last_status = "skipped"
                        self.document.add_page_break()
                    current_feature = line.split(":")[1].lstrip(' ').rstrip()
                    reporter[current_feature] = {}
                    current_scenario = None  # No scenario for the current feature
                    self.document.add_heading(line.rstrip(), 2)
                elif re.match(r"\s*Scenario.*", line):
                    if current_scenario is not None:
                        reporter[current_feature].update({current_scenario: last_status})
                        if last_status == "failed":
                            failed_count += 1
                        elif last_status == "passed":
                            succeed_count += 1
                        test_count += 1
                        last_status = "skipped"
                    current_scenario = line.split(":")[1].lstrip(' ').rstrip()
                    log.debug(current_scenario)
                    self.document.add_heading(line.rstrip(), 3)
                elif re.match(".*passed.*", line):
                    last_status = "passed"
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')
                elif re.match(".*failed.*", line):
                    last_status = "failed"
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')
                else:
                    self.document.add_paragraph(line.rstrip(), style='No Spacing')
        return test_count, succeed_count, failed_count


def insert_text(document, text):
    my_parser = HtmlToDocx()
    md = MarkdownIt()
    md.enable('table')
    my_parser.add_html_to_document(md.render(text), document)